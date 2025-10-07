"""
Comprehensive DOCX Testing Framework
Direct testing of DOCX files with proper RAG integration
"""

import os
import sys
import json
import time
import docx
from datetime import datetime
from typing import Dict, List, Any, Optional
import logging

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Add backend to path
sys.path.append(os.path.join(os.path.dirname(__file__), 'backend'))

# Import backend components
try:
    from backend.rag_engine import RAGEngine
    from backend.scope_validator import ScopeValidator
    from backend.text_chunker import AdvancedTextChunker
    print("✅ Backend components imported successfully")
except ImportError as e:
    print(f"❌ Import error: {e}")
    sys.exit(1)

class ComprehensiveDOCXTester:
    """Comprehensive tester for DOCX files with RAG integration"""
    
    def __init__(self):
        self.results = {
            'test_info': {
                'timestamp': datetime.now().isoformat(),
                'total_tests': 0,
                'passed': 0,
                'failed': 0
            },
            'document_tests': {}
        }
        
        # Test questions from TESTME.md
        self.test_questions = self._load_test_questions()
        
    def _load_test_questions(self) -> Dict[str, Dict]:
        """Load test questions from TESTME.md structure"""
        return {
            'urban_green_spaces': {
                'factual_recall': [
                    "Who designed Central Park in New York?",
                    "In what year was Central Park completed?",
                    "Which cities are named as global leaders in green urban planning?",
                    "How much can trees reduce urban temperatures by?",
                    "What is the United Nations' goal that supports sustainable cities?"
                ],
                'contextual_comparative': [
                    "How did the purpose of green spaces change from the 19th to the 21st century?",
                    "What differences exist between early and modern motivations for preserving green spaces?",
                    "How are modern technologies contributing to green space management?",
                    "What role did the environmental movements of the 1970s and 1980s play in shaping urban policy?"
                ],
                'analytical_reasoning': [
                    'Why is "green gentrification" considered a challenge in urban development?',
                    "How can equitable access to green spaces promote community well-being?",
                    "Why did Frederick Law Olmsted advocate for nature in urban environments?",
                    "How do green spaces contribute to combating climate change in cities?"
                ],
                'summarization': [
                    "Summarize the evolution of urban green spaces from the 19th century to today.",
                    "What are the key benefits and challenges of green spaces in modern cities?",
                    "Explain how sustainability and social equity intersect in the context of urban green spaces."
                ],
                'out_of_scope': [
                    "What year was the Paris Climate Agreement signed?",
                    "Who is the current mayor of Singapore?",
                    "What is the population of New York City?",
                    "Which company created GIS software?"
                ]
            },
            'electric_vehicles': {
                'factual_recall': [
                    "When did the first practical electric car appear?",
                    "What was Tesla's first electric car model, and when was it launched?",
                    "How many miles could the Tesla Roadster travel on a single charge?",
                    "In what year did global EV sales exceed 3 million annually?",
                    "Which United Nations goal supports climate action related to transport?",
                    "What percentage of new car sales are expected to be electric by 2035?"
                ],
                'contextual_comparative': [
                    "How did the popularity of EVs change from the 1920s to the 21st century?",
                    "Compare the technological barriers of early EVs with modern challenges.",
                    "What regions of the world have contributed to accelerating EV adoption?",
                    "How are governments addressing environmental concerns in battery production?"
                ],
                'analytical_reasoning': [
                    "Why did gasoline cars surpass early electric vehicles in the 20th century?",
                    "How do charging innovations aim to make EV use more convenient?",
                    "Why are lithium and nickel central to EV sustainability discussions?",
                    "How does integrating renewable energy with EVs enhance environmental impact?"
                ],
                'summarization': [
                    "Summarize the evolution of electric vehicles from their invention to 2035 projections.",
                    "What are the key environmental benefits and challenges associated with EVs?",
                    "Describe how technology and policy together influence the growth of electric transportation."
                ],
                'out_of_scope': [
                    "What is the range of the Tesla Model S Plaid?",
                    "Who is the current CEO of Tesla?",
                    "What country produces the most crude oil?",
                    "When was the Paris Agreement signed?"
                ]
            }
        }
    
    def extract_docx_content(self, file_path: str) -> Optional[str]:
        """Extract text content from DOCX file"""
        try:
            doc = docx.Document(file_path)
            content_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content_parts.append(text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_texts.append(cell.text.strip())
                    if row_texts:
                        content_parts.append(" | ".join(row_texts))
            
            full_content = '\n\n'.join(content_parts)
            logger.info(f"Extracted {len(full_content)} characters from {file_path}")
            return full_content
            
        except Exception as e:
            logger.error(f"Error extracting DOCX content from {file_path}: {e}")
            return None
    
    def create_rag_with_content(self, content: str, doc_name: str) -> Optional[RAGEngine]:
        """Create RAG engine and add content directly"""
        try:
            # Create RAG engine
            rag_engine = RAGEngine()
            
            # Create text chunks using the chunker
            chunker = AdvancedTextChunker()
            chunks = chunker.chunk_text(content)
            
            logger.info(f"Created {len(chunks)} chunks for {doc_name}")
            
            # Add chunks directly to the RAG engine's vector database
            if hasattr(rag_engine, 'collection') and rag_engine.collection:
                # ChromaDB approach
                chunk_texts = [chunk.content for chunk in chunks]
                chunk_ids = [f"{doc_name}_chunk_{i}" for i in range(len(chunks))]
                chunk_metadatas = [{"source": doc_name, "chunk_id": i} for i in range(len(chunks))]
                
                rag_engine.collection.add(
                    documents=chunk_texts,
                    ids=chunk_ids,
                    metadatas=chunk_metadatas
                )
                
                logger.info(f"Added {len(chunks)} chunks to ChromaDB for {doc_name}")
            else:
                logger.warning(f"ChromaDB not available for {doc_name}, using fallback")
                # Store chunks in a simple format for fallback
                rag_engine._fallback_chunks = chunks
            
            return rag_engine
            
        except Exception as e:
            logger.error(f"Error creating RAG with content: {e}")
            return None
    
    def test_question(self, rag_engine: RAGEngine, question: str, expected_scope: str, category: str) -> Dict:
        """Test a single question with the RAG engine"""
        start_time = time.time()
        
        try:
            # Query the RAG engine
            response = rag_engine.query(question)
            processing_time = time.time() - start_time
            
            # Parse response format
            if isinstance(response, dict):
                answer = response.get('answer', '')
                in_scope = response.get('in_scope', True)
                sources = response.get('sources', [])
            elif isinstance(response, str):
                answer = response
                in_scope = answer != "I don't have information about that in the uploaded documents."
                sources = []
            else:
                answer = str(response)
                in_scope = True
                sources = []
            
            # Evaluate scope correctness
            if expected_scope == 'out_of_scope':
                # Out-of-scope questions should be rejected
                scope_correct = not in_scope or "don't have information" in answer.lower()
            else:
                # In-scope questions should be answered
                scope_correct = in_scope and "don't have information" not in answer.lower()
            
            # Evaluate answer quality for in-scope questions
            quality_score = self._evaluate_answer_quality(answer, expected_scope) if expected_scope != 'out_of_scope' else None
            
            # Determine pass/fail
            if expected_scope == 'out_of_scope':
                passed = scope_correct
            else:
                passed = scope_correct and (quality_score is None or quality_score > 0.3)
            
            return {
                'question': question,
                'category': category,
                'expected_scope': expected_scope,
                'response': {
                    'answer': answer[:300] + "..." if len(answer) > 300 else answer,
                    'in_scope': in_scope,
                    'sources_count': len(sources)
                },
                'evaluation': {
                    'scope_correct': scope_correct,
                    'quality_score': quality_score,
                    'processing_time': processing_time
                },
                'passed': passed,
                'timestamp': datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Error testing question '{question}': {e}")
            return {
                'question': question,
                'category': category,
                'expected_scope': expected_scope,
                'error': str(e),
                'passed': False,
                'processing_time': time.time() - start_time
            }
    
    def _evaluate_answer_quality(self, answer: str, expected_scope: str) -> float:
        """Evaluate the quality of an answer"""
        if not answer or expected_scope == 'out_of_scope':
            return 0.0
        
        answer_lower = answer.lower()
        
        # Check for error indicators
        if any(indicator in answer_lower for indicator in ['error', 'exception', 'failed']):
            return 0.0
        
        # Check for generic non-answers
        if any(phrase in answer_lower for phrase in [
            "don't have information", "can't find", "not available", "outside scope"
        ]):
            return 0.1
        
        # Basic quality scoring
        score = 0.0
        
        # Length indicates substance
        if len(answer) > 50:
            score += 0.4
        
        # Specific information (numbers, proper names, dates)
        import re
        if re.search(r'\b\d{4}\b|\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b|\d+%', answer):
            score += 0.4
        
        # Coherent sentences
        if '. ' in answer or '? ' in answer:
            score += 0.2
        
        return min(score, 1.0)
    
    def test_document(self, doc_key: str, file_path: str) -> Dict:
        """Test a complete document with all question categories"""
        print(f"\n{'='*80}")
        print(f"🧪 Testing: {doc_key.replace('_', ' ').title()}")
        print(f"📄 File: {file_path}")
        print(f"{'='*80}")
        
        # Extract content
        content = self.extract_docx_content(file_path)
        if not content:
            return {'error': f'Failed to extract content from {file_path}'}
        
        print(f"📋 Document content: {len(content)} characters")
        
        # Create RAG engine
        rag_engine = self.create_rag_with_content(content, doc_key)
        if not rag_engine:
            return {'error': f'Failed to create RAG engine for {doc_key}'}
        
        print(f"⚙️ RAG engine initialized successfully")
        
        # Get questions for this document
        questions = self.test_questions.get(doc_key, {})
        if not questions:
            return {'error': f'No test questions found for {doc_key}'}
        
        results = {
            'document_info': {
                'key': doc_key,
                'file_path': file_path,
                'content_length': len(content)
            },
            'category_results': {},
            'summary': {
                'total': 0,
                'passed': 0,
                'failed': 0,
                'pass_rate': 0.0
            }
        }
        
        # Test each category
        for category, question_list in questions.items():
            print(f"\n📊 Testing {category.replace('_', ' ').title()} ({len(question_list)} questions)")\n            \n            category_results = []\n            expected_scope = 'out_of_scope' if category == 'out_of_scope' else 'positive'\n            \n            for i, question in enumerate(question_list, 1):\n                print(f\"  {i:2d}. {question[:70]}{'...' if len(question) > 70 else ''}\")\n                \n                result = self.test_question(rag_engine, question, expected_scope, category)\n                category_results.append(result)\n                \n                # Update counters\n                results['summary']['total'] += 1\n                self.results['test_info']['total_tests'] += 1\n                \n                if result['passed']:\n                    results['summary']['passed'] += 1\n                    self.results['test_info']['passed'] += 1\n                    status = \"✅ PASS\"\n                else:\n                    results['summary']['failed'] += 1\n                    self.results['test_info']['failed'] += 1\n                    status = \"❌ FAIL\"\n                \n                print(f\"      {status}\")\n                \n                # Brief pause to avoid overwhelming\n                time.sleep(0.05)\n            \n            results['category_results'][category] = category_results\n            \n            # Category summary\n            cat_passed = sum(1 for r in category_results if r['passed'])\n            cat_rate = (cat_passed / len(category_results) * 100) if category_results else 0\n            print(f\"  📈 {category.replace('_', ' ').title()}: {cat_rate:.1f}% ({cat_passed}/{len(category_results)})\")\n        \n        # Calculate overall pass rate for document\n        if results['summary']['total'] > 0:\n            results['summary']['pass_rate'] = results['summary']['passed'] / results['summary']['total']\n        \n        return results\n    \n    def run_comprehensive_test(self):\n        \"\"\"Run comprehensive tests on both DOCX files\"\"\"\n        print(\"🚀 Comprehensive DOCX Testing Framework\")\n        print(f\"🕐 Started: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\")\n        print(\"=\"*80)\n        \n        # Test files\n        test_files = {\n            'urban_green_spaces': 'Title_ The Global Impact of Urban Green Spaces.docx',\n            'electric_vehicles': 'Title_ The Rise of Electric Vehicles and the Future of Transportation.docx'\n        }\n        \n        # Test each document\n        for doc_key, file_path in test_files.items():\n            if not os.path.exists(file_path):\n                print(f\"❌ File not found: {file_path}\")\n                continue\n            \n            result = self.test_document(doc_key, file_path)\n            self.results['document_tests'][doc_key] = result\n        \n        # Final summary and analysis\n        self.print_final_summary()\n        self.provide_analysis()\n        self.save_results()\n    \n    def print_final_summary(self):\n        \"\"\"Print comprehensive summary of all tests\"\"\"\n        print(f\"\\n{'='*80}\")\n        print(\"📊 COMPREHENSIVE TEST RESULTS SUMMARY\")\n        print(f\"{'='*80}\")\n        \n        total = self.results['test_info']['total_tests']\n        passed = self.results['test_info']['passed']\n        failed = self.results['test_info']['failed']\n        overall_rate = (passed / total * 100) if total > 0 else 0\n        \n        print(f\"\\n🎯 Overall Performance:\")\n        print(f\"   Total Tests: {total}\")\n        print(f\"   Passed: {passed}\")\n        print(f\"   Failed: {failed}\")\n        print(f\"   Overall Pass Rate: {overall_rate:.1f}%\")\n        \n        print(f\"\\n📚 Document Performance:\")\n        for doc_key, result in self.results['document_tests'].items():\n            if 'error' in result:\n                print(f\"   {doc_key}: ❌ {result['error']}\")\n                continue\n                \n            doc_rate = result['summary']['pass_rate'] * 100\n            doc_name = doc_key.replace('_', ' ').title()\n            print(f\"   {doc_name}: {doc_rate:.1f}% ({result['summary']['passed']}/{result['summary']['total']})\")\n            \n            # Category breakdown\n            best_category = None\n            worst_category = None\n            best_rate = 0\n            worst_rate = 100\n            \n            for category, cat_results in result['category_results'].items():\n                cat_passed = sum(1 for r in cat_results if r['passed'])\n                cat_total = len(cat_results)\n                cat_rate = (cat_passed / cat_total * 100) if cat_total > 0 else 0\n                \n                category_name = category.replace('_', ' ').title()\n                print(f\"     {category_name}: {cat_rate:.1f}% ({cat_passed}/{cat_total})\")\n                \n                if cat_rate > best_rate:\n                    best_rate = cat_rate\n                    best_category = category_name\n                if cat_rate < worst_rate:\n                    worst_rate = cat_rate\n                    worst_category = category_name\n    \n    def provide_analysis(self):\n        \"\"\"Provide detailed analysis and recommendations\"\"\"\n        print(f\"\\n💡 Analysis & Recommendations:\")\n        \n        total = self.results['test_info']['total_tests']\n        passed = self.results['test_info']['passed']\n        overall_rate = (passed / total * 100) if total > 0 else 0\n        \n        if overall_rate >= 80:\n            print(\"   ✅ Excellent performance! System handles DOCX content very well.\")\n        elif overall_rate >= 60:\n            print(\"   ⚠️ Good performance with room for improvement.\")\n        else:\n            print(\"   ❌ Performance needs significant improvement.\")\n        \n        # Analyze category performance\n        category_stats = {}\n        for doc_key, result in self.results['document_tests'].items():\n            if 'category_results' not in result:\n                continue\n                \n            for category, cat_results in result['category_results'].items():\n                if category not in category_stats:\n                    category_stats[category] = {'passed': 0, 'total': 0}\n                \n                cat_passed = sum(1 for r in cat_results if r['passed'])\n                category_stats[category]['passed'] += cat_passed\n                category_stats[category]['total'] += len(cat_results)\n        \n        print(f\"\\n🔍 Category Analysis:\")\n        for category, stats in category_stats.items():\n            if stats['total'] > 0:\n                rate = (stats['passed'] / stats['total'] * 100)\n                category_name = category.replace('_', ' ').title()\n                print(f\"   {category_name}: {rate:.1f}% ({stats['passed']}/{stats['total']})\")\n                \n                if category == 'out_of_scope' and rate < 70:\n                    print(f\"     🔧 Scope validation needs improvement for negative cases\")\n                elif category != 'out_of_scope' and rate < 60:\n                    print(f\"     🔧 Content retrieval and generation needs improvement for {category_name}\")\n    \n    def save_results(self):\n        \"\"\"Save detailed results to JSON file\"\"\"\n        timestamp = datetime.now().strftime(\"%Y%m%d_%H%M%S\")\n        filename = f\"comprehensive_docx_test_results_{timestamp}.json\"\n        \n        try:\n            with open(filename, 'w', encoding='utf-8') as f:\n                json.dump(self.results, f, indent=2, ensure_ascii=False)\n            print(f\"\\n💾 Detailed results saved to: {filename}\")\n        except Exception as e:\n            print(f\"❌ Error saving results: {e}\")\n\ndef main():\n    \"\"\"Main execution function\"\"\"\n    # Check file availability\n    required_files = [\n        'Title_ The Global Impact of Urban Green Spaces.docx',\n        'Title_ The Rise of Electric Vehicles and the Future of Transportation.docx'\n    ]\n    \n    missing_files = [f for f in required_files if not os.path.exists(f)]\n    if missing_files:\n        print(f\"❌ Missing required files: {missing_files}\")\n        return\n    \n    # Run comprehensive tests\n    tester = ComprehensiveDOCXTester()\n    tester.run_comprehensive_test()\n\nif __name__ == \"__main__\":\n    main()