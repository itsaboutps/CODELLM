"""
Document Processor - Extract text from various document formats

Supports PDF, DOCX, TXT, and Markdown files with robust error handling
and text cleaning for optimal RAG performance.
"""

import logging
import io
from typing import Optional, Dict, Any
import streamlit as st

# For PDF processing
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# For DOCX processing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# For text cleaning
import re

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Advanced document processor supporting multiple formats:
    - PDF (using PyPDF2 and pdfplumber for better text extraction)
    - DOCX (using python-docx)
    - TXT (plain text)
    - MD (Markdown)
    """
    
    def __init__(self):
        """Initialize the document processor."""
        self.supported_formats = {
            'pdf': self._extract_pdf,
            'txt': self._extract_txt,
            'docx': self._extract_docx,
            'md': self._extract_markdown
        }
        
        # Log available processors
        available = []
        if PDF_AVAILABLE:
            available.append("PDF")
        if DOCX_AVAILABLE:
            available.append("DOCX")
        available.extend(["TXT", "MD"])
        
        logger.info(f"Document processor initialized. Available formats: {', '.join(available)}")
    
    def process_text(self, text_content: str) -> list:
        """
        Process text content into chunks for RAG system.
        
        Args:
            text_content: Raw text content to process
            
        Returns:
            List of text chunks
        """
        try:
            # Import text chunker
            from .text_chunker import AdvancedTextChunker
            
            # Clean the text first
            cleaned_text = self._clean_text(text_content)
            
            # Initialize chunker and create chunks
            chunker = AdvancedTextChunker()
            chunks = chunker.chunk_text(cleaned_text)
            
            logger.info(f"Processed text into {len(chunks)} chunks")
            return chunks
            
        except Exception as e:
            logger.error(f"Error processing text: {e}")
            return []

    def extract_text(self, uploaded_file) -> Optional[str]:
        """
        Extract text from uploaded file based on its format.
        
        Args:
            uploaded_file: Streamlit uploaded file object
            
        Returns:
            Extracted text content or None if extraction fails
        """
        try:
            # Get file extension
            file_extension = uploaded_file.name.split('.')[-1].lower()
            
            if file_extension not in self.supported_formats:
                logger.error(f"Unsupported file format: {file_extension}")
                return None
            
            # Extract text using appropriate method
            extractor = self.supported_formats[file_extension]
            text_content = extractor(uploaded_file)
            
            if text_content:
                # Clean and normalize the text
                cleaned_text = self._clean_text(text_content)
                logger.info(f"Successfully extracted {len(cleaned_text)} characters from {uploaded_file.name}")
                return cleaned_text
            else:
                logger.error(f"No text content extracted from {uploaded_file.name}")
                return None
                
        except Exception as e:
            logger.error(f"Error extracting text from {uploaded_file.name}: {e}")
            return None
    
    def _extract_pdf(self, uploaded_file) -> Optional[str]:
        """Extract text from PDF file."""
        if not PDF_AVAILABLE:
            logger.error("PDF processing libraries not available. Install PyPDF2 and pdfplumber.")
            return None
        
        try:
            # Reset file pointer
            uploaded_file.seek(0)
            
            # Try pdfplumber first (better for complex layouts)
            try:
                import pdfplumber
                text_content = ""
                with pdfplumber.open(uploaded_file) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += page_text + "\n\n"
                
                if text_content.strip():
                    return text_content
            except Exception as e:
                logger.warning(f"pdfplumber failed, trying PyPDF2: {e}")
            
            # Fallback to PyPDF2
            uploaded_file.seek(0)
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            text_content = ""
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text_content += page_text + "\n\n"
            
            return text_content if text_content.strip() else None
            
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            return None
    
    def _extract_docx(self, uploaded_file) -> Optional[str]:
        """Extract text from DOCX file."""
        if not DOCX_AVAILABLE:
            logger.error("DOCX processing library not available. Install python-docx.")
            return None
        
        try:
            # Reset file pointer
            uploaded_file.seek(0)
            
            # Load document
            doc = Document(uploaded_file)
            
            # Extract text from paragraphs
            text_content = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content += " | ".join(row_text) + "\n"
            
            return text_content if text_content.strip() else None
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            return None
    
    def _extract_txt(self, uploaded_file) -> Optional[str]:
        """Extract text from plain text file."""
        try:
            # Reset file pointer
            uploaded_file.seek(0)
            
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    uploaded_file.seek(0)
                    content = uploaded_file.read()
                    if isinstance(content, bytes):
                        text_content = content.decode(encoding)
                    else:
                        text_content = content
                    
                    return text_content if text_content.strip() else None
                    
                except UnicodeDecodeError:
                    continue
                except Exception as e:
                    logger.warning(f"Error with encoding {encoding}: {e}")
                    continue
            
            logger.error("Could not decode text file with any supported encoding")
            return None
            
        except Exception as e:
            logger.error(f"Error extracting TXT text: {e}")
            return None
    
    def _extract_markdown(self, uploaded_file) -> Optional[str]:
        """Extract text from Markdown file."""
        try:
            # Markdown files are essentially text files
            text_content = self._extract_txt(uploaded_file)
            
            if text_content:
                # Clean up markdown formatting for better chunking
                # Remove markdown headers, links, etc. while preserving content
                cleaned_text = re.sub(r'^#{1,6}\s+', '', text_content, flags=re.MULTILINE)  # Headers
                cleaned_text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', cleaned_text)  # Links
                cleaned_text = re.sub(r'\*\*([^\*]+)\*\*', r'\1', cleaned_text)  # Bold
                cleaned_text = re.sub(r'\*([^\*]+)\*', r'\1', cleaned_text)  # Italic
                cleaned_text = re.sub(r'`([^`]+)`', r'\1', cleaned_text)  # Inline code
                
                return cleaned_text
            
            return None
            
        except Exception as e:
            logger.error(f"Error extracting Markdown text: {e}")
            return None
    
    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text for better processing.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned and normalized text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove excessive newlines but preserve paragraph breaks
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        # Remove special characters that might interfere with processing
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        
        # Normalize quotes
        text = re.sub(r'[""''`]', '"', text)
        text = re.sub(r'[''`]', "'", text)
        
        # Remove excessive punctuation
        text = re.sub(r'[.]{3,}', '...', text)
        text = re.sub(r'[-]{3,}', '---', text)
        
        # Trim and ensure there's content
        text = text.strip()
        
        return text
    
    def get_document_info(self, uploaded_file) -> Dict[str, Any]:
        """
        Get information about the uploaded document.
        
        Args:
            uploaded_file: Streamlit uploaded file object
            
        Returns:
            Dictionary with document information
        """
        try:
            file_extension = uploaded_file.name.split('.')[-1].lower()
            file_size = uploaded_file.size
            
            # Extract text to get character count
            text_content = self.extract_text(uploaded_file)
            char_count = len(text_content) if text_content else 0
            word_count = len(text_content.split()) if text_content else 0
            
            return {
                "filename": uploaded_file.name,
                "format": file_extension.upper(),
                "size_bytes": file_size,
                "size_mb": round(file_size / (1024 * 1024), 2),
                "character_count": char_count,
                "word_count": word_count,
                "supported": file_extension in self.supported_formats,
                "extractable": text_content is not None
            }
            
        except Exception as e:
            logger.error(f"Error getting document info: {e}")
            return {
                "filename": uploaded_file.name,
                "error": str(e)
            }